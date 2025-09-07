from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from typing import Dict, Any

class DocxGenerator:
    def __init__(self):
        self.document = Document()
        self._setup_document()

    def _setup_document(self):
        """Set up document formatting"""
        # Set margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

    def _add_heading(self, text: str, level: int = 1):
        """Add a heading with consistent formatting"""
        heading = self.document.add_heading(text, level=level)
        heading.style.font.name = 'Arial'
        heading.style.font.size = Pt(14 if level == 1 else 12)
        heading.style.font.bold = True
        
        # Add underline for level 1 headings
        if level == 1:
            for run in heading.runs:
                run.underline = True

    def _add_paragraph(self, text: str, bold: bool = False, italic: bool = False):
        """Add a paragraph with consistent formatting"""
        p = self.document.add_paragraph()
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(11)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic

    def _add_bullet_points(self, items: list):
        """Add bullet points with consistent formatting"""
        for item in items:
            p = self.document.add_paragraph(style='List Bullet')
            p.style.font.name = 'Arial'
            p.style.font.size = Pt(11)
            p.add_run(str(item))

    def generate_document(self, data: Dict[str, Any]) -> Document:
        """Generate the standardized resume document"""
        # Add name
        self._add_heading(data.get('name', ''), level=0)
        self.document.add_paragraph(data.get('email', ''))
        
        # Add education
        if data.get('education'):
            self.document.add_paragraph()  # Add spacing
            self._add_heading('EDUCATION', level=1)
            for edu in data['education']:
                self._add_paragraph(f"{edu.get('degree', '')}: {edu.get('institution', '')} : {edu.get('duration', '')}", bold=True)
                if edu.get('details'):
                    self._add_paragraph(edu.get('details', ''))
        
        # Add skills
        if data.get('skills'):
            self.document.add_paragraph()  # Add spacing
            self._add_heading('SKILLS', level=1)
            self._add_bullet_points(data['skills'])
        
        # Add certifications
        if data.get('certifications'):
            self.document.add_paragraph()  # Add spacing
            self._add_heading('CERTIFICATIONS', level=1)
            for cert in data['certifications']:
                if cert.get('link'):
                    # Create hyperlink (simplified representation)
                    cert_text = f"{cert.get('name', '')} ({cert.get('issuer', '')})"
                    self._add_paragraph(cert_text)
                else:
                    self._add_paragraph(f"{cert.get('name', '')} ({cert.get('issuer', '')})")
        
        # Add professional experience
        if data.get('experience'):
            self.document.add_paragraph()  # Add spacing
            self._add_heading('PROFESSIONAL EXPERIENCE', level=1)
            
            for exp in data['experience']:
                # Add role and company with duration
                role_text = f"{exp.get('role', '')} at {exp.get('company', '')} ({exp.get('duration', '')})"
                self._add_paragraph(role_text, bold=True)
                
                # Add responsibilities
                if exp.get('responsibilities'):
                    self._add_bullet_points(exp['responsibilities'])
                
                self.document.add_paragraph()  # Add spacing between experiences
        
        # Add achievements
        if data.get('achievements'):
            self.document.add_paragraph()  # Add spacing
            self._add_heading('ACHIEVEMENTS', level=1)
            self._add_bullet_points(data['achievements'])
        
        return self.document

    def save_document(self, data: Dict[str, Any], filename: str):
        """Generate and save the document"""
        doc = self.generate_document(data)
        doc.save(filename)